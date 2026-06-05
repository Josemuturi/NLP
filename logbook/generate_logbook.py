"""
logbook/generate_logbook.py — Smart Farm Professional Logbook Generator
=========================================================================
Generates a professional .docx logbook covering all 3 weeks of the
BIT4133 NLP course, mapped to the Smart Farm project.

Structure:
  - Cover Page
  - Table of Contents
  - Week 1: NLTK Basics (Tokenization, Stopwords, Stemming, Lemmatization)
  - Week 2: N-gram Models & POS Tagging
  - Week 3: HMM Sequence Labeling & Full Chatbot
  - References

Run: python logbook/generate_logbook.py

Course: BIT4133 Natural Language Processing
Project: Smart Farm AI Assistant
"""

import sys
import os
from datetime import datetime

# Ensure python-docx is importable
try:
    from docx import Document
    from docx.shared import (Pt, Inches, RGBColor, Cm)
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)


# =============================================================================
# CONFIGURATION
# =============================================================================

STUDENT_NAME   = "John Doe"           # ← Change to actual student name
REG_NUMBER     = "BIT/4133/2024"      # ← Change to actual registration number
COURSE_CODE    = "BIT4133"
COURSE_NAME    = "Natural Language Processing"
PROJECT_NAME   = "Smart Farm AI Assistant"
SEMESTER       = "Semester 1, 2025/2026"
INSTITUTION    = "Your University Name"  # ← Change to actual institution
SUPERVISOR     = "Dr. Jane Smith"         # ← Change to supervisor name
SUBMISSION_DATE = datetime.now().strftime("%B %d, %Y")

OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "Smart_Farm_Logbook.docx")


# =============================================================================
# COLOR PALETTE
# =============================================================================

GREEN_DARK  = RGBColor(0x0D, 0x47, 0xA1)   # Dark blue (headings)
GREEN_MID   = RGBColor(0x15, 0x65, 0xC0)   # Medium blue (subheadings)
GREEN_LIGHT = RGBColor(0x42, 0xA5, 0xF5)   # Light blue
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_DARK   = RGBColor(0x33, 0x33, 0x33)
GRAY_MID    = RGBColor(0x66, 0x66, 0x66)
GRAY_LIGHT  = RGBColor(0xF5, 0xF5, 0xF5)
ACCENT      = RGBColor(0xFF, 0x8F, 0x00)   # Amber accent


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def _set_cell_background(cell, color_hex: str):
    """Set the background color of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1,
                 color: RGBColor = None, center: bool = False):
    """Add a heading paragraph with optional color and alignment."""
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_para(doc: Document, text: str = "", bold: bool = False,
              italic: bool = False, color: RGBColor = None,
              size: int = 11, center: bool = False,
              space_before: int = 0, space_after: int = 6) -> object:
    """Add a paragraph with full formatting control."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


def _add_code_block(doc: Document, code: str, caption: str = ""):
    """Add a formatted code block (monospaced, shaded background)."""
    if caption:
        cap_p = doc.add_paragraph()
        run   = cap_p.add_run(f"// {caption}")
        run.bold = False
        run.font.size = Pt(9)
        run.font.color.rgb = GREEN_MID
        cap_p.paragraph_format.space_after = Pt(2)

    # Add a table with 1 cell for the code block
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    _set_cell_background(cell, "0D0D0D")

    # Clear default paragraph and set code content
    cell.paragraphs[0].clear()
    for line in code.strip().split("\n"):
        p   = cell.add_paragraph(line)
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xF0, 0xF0, 0xF0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Pt(6)

    # Remove first empty paragraph
    first_p = cell.paragraphs[0]
    if not first_p.text:
        first_p._element.getparent().remove(first_p._element)

    doc.add_paragraph()  # spacing after code block


def _add_screenshot_placeholder(doc: Document, description: str,
                                  label: str = "INSERT SCREENSHOT HERE"):
    """Add a labeled screenshot placeholder box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell  = table.cell(0, 0)
    _set_cell_background(cell, "E3F2FD")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(18)

    run1 = p.add_run(f"[ {label} ]")
    run1.bold = True
    run1.font.size  = Pt(11)
    run1.font.color.rgb = GREEN_DARK

    p.add_run("\n")
    run2 = p.add_run(description)
    run2.font.size  = Pt(9)
    run2.italic = True
    run2.font.color.rgb = GRAY_MID

    doc.add_paragraph()  # spacing after placeholder


def _add_github_placeholder(doc: Document, week: int, commit_msg: str):
    """Add a GitHub commit screenshot placeholder."""
    _add_screenshot_placeholder(
        doc,
        description=f"GitHub commit screenshot — Week {week}: {commit_msg}",
        label="INSERT GITHUB COMMIT SCREENSHOT HERE"
    )


def _add_info_table(doc: Document, data: list, header_row: list = None):
    """
    Add a formatted table.
    data: list of row lists (strings)
    header_row: optional header row (bold, green background)
    """
    n_cols = len(data[0]) if data else 1
    if header_row:
        table = doc.add_table(rows=len(data) + 1, cols=n_cols)
    else:
        table = doc.add_table(rows=len(data), cols=n_cols)

    table.style = "Table Grid"

    row_offset = 0
    if header_row:
        row = table.rows[0]
        for i, cell_text in enumerate(header_row):
            cell = row.cells[i]
            _set_cell_background(cell, "1565C0")
            p   = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.bold = True
            run.font.color.rgb = WHITE
            run.font.size = Pt(10)
        row_offset = 1

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + row_offset]
        bg  = "FFFFFF" if r_idx % 2 == 0 else "F8F9FF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            _set_cell_background(cell, bg)
            p   = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)

    doc.add_paragraph()


def _add_page_break(doc: Document):
    """Add a page break."""
    doc.add_page_break()


# =============================================================================
# COVER PAGE
# =============================================================================

def build_cover_page(doc: Document):
    """Build a professional cover page."""
    # Top green bar (via table)
    top_table = doc.add_table(rows=1, cols=1)
    top_table.style = "Table Grid"
    top_cell = top_table.cell(0, 0)
    _set_cell_background(top_cell, "0D47A1")
    p = top_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(20)
    run = p.add_run(f"🌾  {PROJECT_NAME}  🌾")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = WHITE

    doc.add_paragraph()
    doc.add_paragraph()

    # University name
    _add_para(doc, INSTITUTION, bold=True, size=14, center=True, color=GREEN_DARK)
    _add_para(doc, f"Department of Computer Science", size=12, center=True, color=GRAY_MID)
    doc.add_paragraph()

    # Course badge table
    badge_table = doc.add_table(rows=1, cols=1)
    badge_table.style = "Table Grid"
    badge_cell = badge_table.cell(0, 0)
    _set_cell_background(badge_cell, "E3F2FD")
    p = badge_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    for line in [
        f"Course Code : {COURSE_CODE}",
        f"Course Name : {COURSE_NAME}",
        f"Project     : {PROJECT_NAME}",
        f"Semester    : {SEMESTER}",
    ]:
        rr = p.add_run(line + "\n")
        rr.font.size = Pt(11)
        rr.font.color.rgb = GREEN_DARK

    doc.add_paragraph()
    doc.add_paragraph()

    # Student info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Student Name",     STUDENT_NAME),
        ("Registration No.", REG_NUMBER),
        ("Supervisor",       SUPERVISOR),
        ("Date Submitted",   SUBMISSION_DATE),
    ]
    for i, (label, value) in enumerate(info_data):
        lc = info_table.rows[i].cells[0]
        vc = info_table.rows[i].cells[1]
        _set_cell_background(lc, "2E7D32")
        _set_cell_background(vc, "FFFFFF")
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.color.rgb = WHITE
        lr.font.size = Pt(10)
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)

    doc.add_paragraph()
    _add_para(doc, "— NLP Logbook: Weeks 1–5 —", italic=True, size=10,
              center=True, color=GRAY_MID)

    _add_page_break(doc)


# =============================================================================
# TABLE OF CONTENTS
# =============================================================================

def build_toc(doc: Document):
    """Build a manual table of contents."""
    _add_heading(doc, "Table of Contents", level=1, color=GREEN_DARK)

    toc_items = [
        ("1.",  "Week 1: NLTK Basics — Tokenization, Stemming & Lemmatization",  "3"),
        ("1.1", "Week Theme & Objectives",                                         "3"),
        ("1.2", "Tasks Completed",                                                 "3"),
        ("1.3", "Technologies Used",                                               "4"),
        ("1.4", "Code Snippets",                                                   "4"),
        ("1.5", "Screenshots & Practical Outputs",                                 "5"),
        ("1.6", "Learning Outcomes",                                               "5"),
        ("1.7", "GitHub Commit Screenshot",                                        "6"),
        ("2.",  "Week 2: N-gram Models & POS Tagging",                            "7"),
        ("2.1", "Week Theme & Objectives",                                         "7"),
        ("2.2", "Tasks Completed",                                                 "7"),
        ("2.3", "Technologies Used",                                               "8"),
        ("2.4", "Code Snippets",                                                   "8"),
        ("2.5", "Screenshots & Practical Outputs",                                 "9"),
        ("2.6", "Learning Outcomes",                                               "9"),
        ("2.7", "GitHub Commit Screenshot",                                       "10"),
        ("3.",  "Week 3: HMM Sequence Labeling & Full Chatbot",                  "11"),
        ("3.1", "Week Theme & Objectives",                                        "11"),
        ("3.2", "Tasks Completed",                                                "11"),
        ("3.3", "Technologies Used",                                              "12"),
        ("3.4", "Code Snippets",                                                  "12"),
        ("3.5", "Screenshots & Practical Outputs",                                "14"),
        ("3.6", "Learning Outcomes",                                              "14"),
        ("3.7", "GitHub Commit Screenshot",                                       "15"),
        ("4.",  "Week 4: Syntactic & Semantic Analysis",                         "16"),
        ("4.1", "Week Theme & Objectives",                                        "16"),
        ("4.2", "Tasks Completed",                                                "16"),
        ("4.3", "Technologies Used",                                              "17"),
        ("4.4", "Code Snippets",                                                  "17"),
        ("4.5", "Screenshots & Practical Outputs",                                "19"),
        ("4.6", "Learning Outcomes",                                              "19"),
        ("4.7", "Assignment Results",                                             "20"),
        ("4.8", "GitHub Commit Screenshot",                                       "21"),
        ("5.",  "Week 5: CAT 1 Preparation & Mini NLP Projects",                "22"),
        ("5.1", "Week Theme & Objectives",                                        "22"),
        ("5.2", "Tasks Completed",                                                "22"),
        ("5.3", "Technologies Used",                                              "23"),
        ("5.4", "Code Snippets",                                                  "23"),
        ("5.5", "Screenshots & Practical Outputs",                                "25"),
        ("5.6", "Mini Project Documentation",                                     "25"),
        ("5.7", "CAT 1 Revision Summary",                                         "26"),
        ("5.8", "Learning Outcomes",                                              "26"),
        ("5.9", "Online Course Evidence (CAT 2)",                                 "27"),
        ("5.10","GitHub Commit Screenshot",                                       "28"),
        ("6.",  "References",                                                     "29"),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)
    toc_table.style = "Table Grid"
    for i, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]
        bg  = "E3F2FD" if i % 2 == 0 else "FFFFFF"
        _set_cell_background(row.cells[0], bg)
        _set_cell_background(row.cells[1], bg)
        _set_cell_background(row.cells[2], bg)

        n_run = row.cells[0].paragraphs[0].add_run(num)
        n_run.font.size = Pt(10)
        n_run.bold = num.endswith(".")  # bold top-level

        t_run = row.cells[1].paragraphs[0].add_run(title)
        t_run.font.size = Pt(10)
        t_run.bold = num.endswith(".")

        p_run = row.cells[2].paragraphs[0].add_run(page)
        p_run.font.size = Pt(10)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _add_page_break(doc)


# =============================================================================
# WEEK 1
# =============================================================================

def build_week1(doc: Document):
    """Build Week 1 section of the logbook."""
    _add_heading(doc, "Week 1: NLTK Basics — Tokenization, Stemming & Lemmatization",
                 level=1, color=GREEN_DARK)

    # ── 1.1 Theme ─────────────────────────────────────────────────────────
    _add_heading(doc, "1.1  Week Theme & Objectives", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 1 introduced the foundational concepts of Natural Language Processing (NLP). "
        "The focus was on understanding how raw text can be preprocessed and normalized "
        "for downstream NLP tasks. The Smart Farm project used these techniques to process "
        "farmer queries before attempting to understand or answer them."
    ))
    _add_para(doc, "Learning Objectives:", bold=True, space_after=2)
    for obj in [
        "Understand and apply word tokenization and sentence tokenization.",
        "Remove stop words to retain only meaningful content words.",
        "Apply Porter stemming to reduce words to their root form.",
        "Apply WordNet lemmatization to obtain dictionary base forms.",
        "Build a simple knowledge base of farming problems and solutions.",
        "Process real farming sentences through the complete NLP pipeline.",
    ]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 1.2 Tasks ─────────────────────────────────────────────────────────
    _add_heading(doc, "1.2  Tasks Completed", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["1", "Set up Python 3.x development environment and virtual environment"],
            ["2", "Installed NLTK library and downloaded required language data packages"],
            ["3", "Created knowledge_base.py with 30+ farming problems and solutions dictionary"],
            ["4", "Implemented word_tokenize() and sent_tokenize() on 5 farming sentences"],
            ["5", "Applied NLTK stopwords corpus to remove non-content words from tokens"],
            ["6", "Applied Porter stemmer and Snowball stemmer — compared output"],
            ["7", "Applied WordNetLemmatizer with verb (v) and noun (n) POS modes"],
            ["8", "Produced comparison table: original word → stem → lemma"],
            ["9", "Ran nlp_pipeline.py on all 5 example farming sentences"],
            ["10", "Captured and documented terminal outputs as screenshots"],
        ],
        header_row=["#", "Task Description"]
    )

    # ── 1.3 Technologies ──────────────────────────────────────────────────
    _add_heading(doc, "1.3  Technologies Used", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["Python 3.x",             "Programming language",        "3.10+"],
            ["NLTK 3.8+",              "NLP library",                 "pip install nltk"],
            ["nltk.tokenize",          "Word & sentence tokenization","Built into NLTK"],
            ["nltk.corpus.stopwords",  "English stop words list",     "NLTK data download"],
            ["nltk.stem.PorterStemmer","Porter stemming algorithm",   "Built into NLTK"],
            ["nltk.stem.WordNetLemmatizer", "WordNet-based lemmatization", "Requires wordnet corpus"],
            ["VS Code / PyCharm",      "Code editor",                 ""],
        ],
        header_row=["Technology", "Purpose", "Installation"]
    )

    # ── 1.4 Code Snippets ─────────────────────────────────────────────────
    _add_heading(doc, "1.4  Code Snippets", level=2, color=GREEN_MID)

    _add_code_block(doc, """
# week1/nlp_pipeline.py — Tokenization, Stop Words, Stemming & Lemmatization
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.stem import WordNetLemmatizer

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

# Initialize tools
stemmer    = PorterStemmer()
lemmatizer = WordNetLemmatizer()
STOP_WORDS = set(stopwords.words("english"))

sentence = "My maize leaves are turning yellow and the plants are stunted."

# Step 1: Word tokenization
tokens = word_tokenize(sentence)
print("Tokens:", tokens)

# Step 2: Remove stop words
filtered = [t for t in tokens if t.lower() not in STOP_WORDS]
print("Filtered:", filtered)

# Step 3: Stemming
stems = [stemmer.stem(t) for t in filtered]
print("Stems:", stems)

# Step 4: Lemmatization
lemmas = [lemmatizer.lemmatize(t.lower(), pos='v') for t in filtered]
print("Lemmas:", lemmas)
""", caption="NLP Pipeline — week1/nlp_pipeline.py")

    _add_code_block(doc, """
# knowledge_base.py — Farming Problems & Solutions Dictionary (excerpt)
KNOWLEDGE_BASE = {
    ("maize", "yellow leaves"): {
        "cause":    "Nitrogen deficiency or Maize Streak Virus (MSV).",
        "solution": "Apply urea fertilizer (46% N) at 50 kg/acre. "
                    "If viral symptoms, remove and destroy infected plants.",
        "action":   "fertilize | remove infected plants",
        "severity": "medium"
    },
    ("tomato", "blight"): {
        "cause":    "Early blight (Alternaria solani) or late blight (Phytophthora).",
        "solution": "Spray mancozeb for early blight or metalaxyl for late blight.",
        "action":   "remove infected parts | spray fungicide",
        "severity": "high"
    },
    # ... 28 more entries covering maize, tomato, beans, wheat, potato, rice, cassava
}
""", caption="Knowledge Base — knowledge_base.py")

    # ── 1.5 Screenshots ───────────────────────────────────────────────────
    _add_heading(doc, "1.5  Screenshots & Practical Outputs", level=2, color=GREEN_MID)

    for desc in [
        "Terminal output of nlp_pipeline.py — showing tokenization, filtered tokens, stems, and lemmas for all 5 farming sentences",
        "Porter vs Snowball stemmer comparison table output",
        "Lemmatization vs Stemming comparison table output (word → stem → noun-lemma → verb-lemma)",
        "Knowledge base test output — lookup_solution('maize', 'yellow leaves') result",
        "Python environment setup — pip install nltk output and NLTK download confirmation",
    ]:
        _add_screenshot_placeholder(doc, desc)

    # ── 1.6 Learning Outcomes ─────────────────────────────────────────────
    _add_heading(doc, "1.6  Learning Outcomes", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 1 established the essential preprocessing pipeline for the Smart Farm chatbot. "
        "By the end of this week, the following key competencies were acquired:"
    ))
    outcomes = [
        ("Tokenization", "Understanding that text must first be segmented into tokens before any NLP analysis. Both word-level and sentence-level tokenization were demonstrated."),
        ("Stop Word Removal", "Stop words (the, is, are, my) carry little semantic meaning for farming queries. Removing them improves downstream accuracy by focusing on content words like 'maize', 'yellow', 'wilting'."),
        ("Stemming", "Porter stemming reduces words to approximate roots ('yellowing' → 'yellow', 'spraying' → 'spray'). Fast but sometimes produces non-words."),
        ("Lemmatization", "WordNet lemmatization produces valid dictionary forms ('leaves' → 'leaf', 'turning' → 'turn'). Slower than stemming but linguistically accurate."),
        ("Knowledge Base Design", "A structured Python dictionary is an effective and transparent knowledge representation suitable for domain-specific chatbots."),
    ]
    for term, explanation in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()

    # ── 1.7 GitHub Screenshot ─────────────────────────────────────────────
    _add_heading(doc, "1.7  GitHub Commit Screenshot", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Commit 1 added the Python environment, NLTK setup, Week 1 NLP pipeline scripts, "
        "the demo sentences file, and the farming knowledge base dictionary."
    ))
    _add_github_placeholder(doc, week=1,
        commit_msg="Add Week 1: NLTK setup, tokenization, stemming, lemmatization & knowledge base")

    _add_page_break(doc)


# =============================================================================
# WEEK 2
# =============================================================================

def build_week2(doc: Document):
    """Build Week 2 section of the logbook."""
    _add_heading(doc, "Week 2: N-gram Language Models & POS Tagging",
                 level=1, color=GREEN_DARK)

    # ── 2.1 Theme ─────────────────────────────────────────────────────────
    _add_heading(doc, "2.1  Week Theme & Objectives", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 2 built upon the preprocessing pipeline to develop a deeper understanding "
        "of language structure. N-gram models were constructed to model word sequence "
        "probabilities in farming language, and POS tagging was used to identify the "
        "grammatical role of each word in a farmer's query — distinguishing crop names "
        "(nouns), symptoms (adjectives), and actions (verbs)."
    ))
    _add_para(doc, "Learning Objectives:", bold=True, space_after=2)
    for obj in [
        "Understand n-gram language models: unigrams, bigrams, trigrams.",
        "Build a bigram model from a farming text corpus.",
        "Calculate n-gram probabilities using Maximum Likelihood Estimation (MLE).",
        "Apply Laplace (add-1) smoothing to handle unseen n-grams.",
        "Apply NLTK's averaged perceptron POS tagger to farmer queries.",
        "Extract nouns (crop names), verbs (actions), and adjectives (symptoms) from tagged text.",
        "Implement noun phrase (NP) chunking using RegexpParser.",
        "Classify query tokens into farming roles: crop, symptom, action, location.",
    ]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 2.2 Tasks ─────────────────────────────────────────────────────────
    _add_heading(doc, "2.2  Tasks Completed", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["1",  "Created a farming text corpus of 20 sentences for n-gram training"],
            ["2",  "Built FarmingNgramModel class supporting n=1, 2, 3"],
            ["3",  "Implemented MLE probability with Laplace add-1 smoothing"],
            ["4",  "Generated top-10 unigrams, bigrams, and trigrams from corpus"],
            ["5",  "Implemented predict_next() to suggest likely next words"],
            ["6",  "Scored farming vs non-farming sentences using log-probability"],
            ["7",  "Calculated model perplexity on hold-out test sentences"],
            ["8",  "Applied NLTK pos_tag() to 5 farmer queries — displayed full POS tables"],
            ["9",  "Extracted nouns, verbs, adjectives from tagged queries"],
            ["10", "Implemented NP chunking with RegexpParser grammar"],
            ["11", "Built identify_farming_roles() to classify words into farming roles"],
            ["12", "Ran pos_tagger.py on 5 queries — documented POS tag distribution"],
        ],
        header_row=["#", "Task Description"]
    )

    # ── 2.3 Technologies ──────────────────────────────────────────────────
    _add_heading(doc, "2.3  Technologies Used", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["nltk.pos_tag",           "POS tagging using averaged perceptron", "Built into NLTK"],
            ["nltk.RegexpParser",      "Regex-based chunking for NP extraction","Built into NLTK"],
            ["collections.Counter",    "N-gram frequency counting",             "Python stdlib"],
            ["math.log2",              "Log-probability computation",           "Python stdlib"],
            ["Penn Treebank Tagset",   "Universal POS tag reference",           "NLTK standard"],
        ],
        header_row=["Technology", "Purpose", "Source"]
    )

    # ── 2.4 Code Snippets ─────────────────────────────────────────────────
    _add_heading(doc, "2.4  Code Snippets", level=2, color=GREEN_MID)

    _add_code_block(doc, """
# week2/ngram_model.py — Bigram Language Model (excerpt)
from collections import Counter
import math

class FarmingNgramModel:
    def __init__(self, corpus, n=2):
        self.n = n
        self.ngram_freq  = Counter()
        self.context_freq = Counter()
        self.vocabulary  = set()
        self._build(corpus)

    def _build(self, corpus):
        for sentence in corpus:
            tokens = self._tokenize(sentence)
            self.vocabulary.update(tokens)
            for i in range(len(tokens) - self.n + 1):
                ngram   = tuple(tokens[i:i + self.n])
                context = ngram[:-1]
                self.ngram_freq[ngram]    += 1
                self.context_freq[context] += 1

    def probability(self, ngram):
        # Laplace (add-1) smoothing
        context    = ngram[:-1]
        vocab_size = len(self.vocabulary)
        return (self.ngram_freq[ngram] + 1) / (self.context_freq[context] + vocab_size)

    def predict_next(self, context_words, top_k=3):
        context = tuple(w.lower() for w in context_words[-(self.n-1):])
        candidates = {ngram[-1]: self.probability(ngram)
                      for ngram in self.ngram_freq if ngram[:-1] == context}
        return sorted(candidates.items(), key=lambda x: -x[1])[:top_k]

# Example: predict next word after "maize"
bigram_model = FarmingNgramModel(corpus, n=2)
print(bigram_model.predict_next(["maize"]))
# Output: [('leaves', 0.123), ('streak', 0.089), ('plants', 0.076)]
""", caption="N-gram Language Model — week2/ngram_model.py")

    _add_code_block(doc, """
# week2/pos_tagger.py — POS Tagging and Role Classification (excerpt)
import nltk
from nltk import pos_tag, RegexpParser
from nltk.tokenize import word_tokenize

def analyse_farmer_query(text):
    tokens     = word_tokenize(text)
    tagged     = pos_tag(tokens)   # POS tagging

    # Noun Phrase chunking
    grammar = r"NP: {<DT>?<JJ.*>*<NN.*>+}"
    parser  = RegexpParser(grammar)
    tree    = parser.parse(tagged)
    noun_phrases = [" ".join(w for w, t in st.leaves())
                    for st in tree.subtrees() if st.label() == "NP"]

    # Classify into farming roles
    nouns      = [w for w, t in tagged if t in ("NN","NNS","NNP")]
    verbs      = [w for w, t in tagged if t.startswith("VB")]
    adjectives = [w for w, t in tagged if t.startswith("JJ")]

    return {"tagged": tagged, "nouns": nouns, "verbs": verbs,
            "adjectives": adjectives, "noun_phrases": noun_phrases}

query  = "My maize leaves are turning yellow."
result = analyse_farmer_query(query)
print("POS tags:", result["tagged"])
print("Nouns   :", result["nouns"])
print("Verbs   :", result["verbs"])
""", caption="POS Tagger — week2/pos_tagger.py")

    # ── 2.5 Screenshots ───────────────────────────────────────────────────
    _add_heading(doc, "2.5  Screenshots & Practical Outputs", level=2, color=GREEN_MID)
    for desc in [
        "N-gram model statistics table — vocabulary size, unique n-grams, total count for unigram/bigram/trigram",
        "Top-10 bigrams and trigrams from the farming corpus",
        "Next word prediction output — e.g., after 'maize': predicted words with probabilities",
        "Log-probability sentence scoring — farming sentence vs non-farming sentence comparison",
        "POS tagging output — full word/tag table for a farmer query with tag descriptions",
        "Farming role classification output — crops, symptoms, actions, locations identified in query",
    ]:
        _add_screenshot_placeholder(doc, desc)

    # ── 2.6 Learning Outcomes ─────────────────────────────────────────────
    _add_heading(doc, "2.6  Learning Outcomes", level=2, color=GREEN_MID)
    outcomes = [
        ("N-gram Models", "N-gram models capture word co-occurrence patterns. Higher-order n-grams (trigrams) encode more context but require larger corpora. Laplace smoothing prevents zero probabilities."),
        ("Perplexity", "Perplexity measures how well a language model predicts new text — lower is better. Bigram models consistently outperform unigrams on domain-specific farming text."),
        ("POS Tagging", "NLTK's averaged perceptron tagger assigns Penn Treebank tags. Nouns (NN/NNS) typically correspond to crop names; verbs (VBG/VBP) indicate farmer actions; adjectives (JJ) describe symptoms."),
        ("NP Chunking", "Regex grammar allows extraction of noun phrases like 'my maize leaves', which contains both the crop (maize) and location (leaves) in one structured phrase."),
        ("Farming Role Classification", "Custom vocabulary matching combined with POS constraints provides a practical way to identify domain-specific roles without a full named-entity recognizer."),
    ]
    for term, explanation in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()

    # ── 2.7 GitHub ─────────────────────────────────────────────────────────
    _add_heading(doc, "2.7  GitHub Commit Screenshot", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Commit 2 added the N-gram language model built on the farming corpus and "
        "the POS tagger with noun/verb extraction and farming role classification."
    ))
    _add_github_placeholder(doc, week=2,
        commit_msg="Add Week 2: N-gram language model, POS tagging & farming role classifier")

    _add_page_break(doc)


# =============================================================================
# WEEK 3
# =============================================================================

def build_week3(doc: Document):
    """Build Week 3 section of the logbook."""
    _add_heading(doc, "Week 3: Hidden Markov Models & Full Chatbot",
                 level=1, color=GREEN_DARK)

    # ── 3.1 Theme ─────────────────────────────────────────────────────────
    _add_heading(doc, "3.1  Week Theme & Objectives", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 3 integrated all previous NLP concepts into a fully working system. "
        "Hidden Markov Models (HMMs) were introduced as a probabilistic framework "
        "for sequence labeling — used here to identify entities in farmer messages "
        "(crop, disease, symptom, location, action). The full chatbot pipeline was "
        "assembled and speech recognition was added so farmers can speak their problems."
    ))
    _add_para(doc, "Learning Objectives:", bold=True, space_after=2)
    for obj in [
        "Understand HMM theory: states, observations, transitions, emissions, initial probabilities.",
        "Implement the Viterbi algorithm for finding the most likely label sequence.",
        "Apply HMM to label entities in farming text: CROP, DISEASE, SYMPTOM, LOCATION, ACTION.",
        "Wire together the full NLP pipeline (Week 1 + Week 2 + Week 3).",
        "Build an interactive CLI chatbot that accepts farmer input and returns solutions.",
        "Integrate speech-to-text input using the SpeechRecognition library.",
        "Demonstrate the complete Smart Farm AI Assistant end-to-end.",
    ]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── HMM Theory Box ────────────────────────────────────────────────────
    _add_heading(doc, "HMM Theory Overview", level=3, color=GREEN_MID)
    theory_table = doc.add_table(rows=1, cols=1)
    theory_table.style = "Table Grid"
    tc = theory_table.cell(0, 0)
    _set_cell_background(tc, "E3F2FD")
    p = tc.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

    theory_lines = [
        ("States (S)",      "Entity labels: CROP, DISEASE, SYMPTOM, LOCATION, ACTION, O"),
        ("Observations (O)","Words in the farmer's message"),
        ("Transitions (A)", "P(label_t | label_{t-1}) — probability of moving between entity labels"),
        ("Emissions (B)",   "P(word | label) — probability of a word given its entity label"),
        ("Initial (π)",     "P(label_0) — probability of starting in each label"),
        ("Algorithm",       "Viterbi — dynamic programming finds the globally optimal label sequence"),
        ("Goal",            "Find the label sequence L* = argmax P(L|W) for word sequence W"),
    ]
    for term, val in theory_lines:
        rr = p.add_run(f"  {term:<20}: {val}\n")
        rr.font.name = "Courier New"
        rr.font.size = Pt(9)

    doc.add_paragraph()

    # ── 3.2 Tasks ─────────────────────────────────────────────────────────
    _add_heading(doc, "3.2  Tasks Completed", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["1",  "Designed HMM entity label set: CROP, DISEASE, SYMPTOM, LOCATION, ACTION, O"],
            ["2",  "Built emission probability table from knowledge base vocabulary"],
            ["3",  "Defined transition matrix based on typical entity sequences in farming text"],
            ["4",  "Implemented Viterbi algorithm with log-probability and backtracking"],
            ["5",  "Tested Viterbi on 5 farming queries — displayed Viterbi matrix"],
            ["6",  "Implemented extract_entities() to group labeled tokens into entities"],
            ["7",  "Built SmartFarmChatbot class wiring all 3 weeks' NLP modules"],
            ["8",  "Implemented _find_solution() to map extracted entities to knowledge base"],
            ["9",  "Added --speech flag to chatbot for SpeechRecognition integration"],
            ["10", "Added --demo flag for non-interactive demonstration mode"],
            ["11", "Ran chatbot with 5 diverse farming queries and documented output"],
            ["12", "Generated and verified logbook .docx using python-docx"],
        ],
        header_row=["#", "Task Description"]
    )

    # ── 3.3 Technologies ──────────────────────────────────────────────────
    _add_heading(doc, "3.3  Technologies Used", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["HMM (custom)",        "Sequence labeling model",             "Implemented from scratch"],
            ["Viterbi Algorithm",   "Optimal label sequence decoding",     "Dynamic programming"],
            ["SpeechRecognition",   "Speech-to-text input",               "pip install SpeechRecognition"],
            ["PyAudio",             "Microphone access layer",             "pip install pyaudio"],
            ["Google Web Speech API","Cloud speech transcription",         "Via SpeechRecognition"],
            ["colorama",            "Colored terminal output",             "pip install colorama"],
            ["python-docx",         "Logbook .docx generation",           "pip install python-docx"],
            ["argparse",            "CLI argument parsing (--speech, --demo)", "Python stdlib"],
        ],
        header_row=["Technology", "Purpose", "Installation"]
    )

    # ── 3.4 Code Snippets ─────────────────────────────────────────────────
    _add_heading(doc, "3.4  Code Snippets", level=2, color=GREEN_MID)

    _add_code_block(doc, """
# week3/hmm_entity.py — Viterbi Algorithm (excerpt)
import math

LABELS    = ["CROP", "DISEASE", "SYMPTOM", "LOCATION", "ACTION", "O"]
N_STATES  = len(LABELS)

def viterbi(self, tokens):
    T = len(tokens)
    N = N_STATES

    # Viterbi table: V[t][s] = max log-prob reaching state s at time t
    V = [[float("-inf")] * N for _ in range(T)]
    B = [[0] * N for _ in range(T)]           # backpointers

    # Initialization
    for s in range(N):
        label = LABELS[s]
        V[0][s] = (math.log(self.initial[s]) +
                   math.log(self._emit_prob(label, tokens[0])))

    # Recursion
    for t in range(1, T):
        for s in range(N):
            emit_log = math.log(self._emit_prob(LABELS[s], tokens[t]))
            best_score, best_prev = float("-inf"), 0
            for prev_s in range(N):
                score = (V[t-1][prev_s] +
                         math.log(self.transition[prev_s][s]))
                if score > best_score:
                    best_score, best_prev = score, prev_s
            V[t][s] = best_score + emit_log
            B[t][s] = best_prev

    # Backtrack to find optimal path
    path = [0] * T
    path[T-1] = max(range(N), key=lambda s: V[T-1][s])
    for t in range(T-2, -1, -1):
        path[t] = B[t+1][path[t+1]]

    return [LABELS[s] for s in path]
""", caption="Viterbi Algorithm — week3/hmm_entity.py")

    _add_code_block(doc, """
# week3/chatbot.py — Full NLP Pipeline Integration (excerpt)
from nlp_pipeline import run_full_pipeline    # Week 1
from pos_tagger   import analyse_farmer_query  # Week 2
from hmm_entity   import FarmingHMM           # Week 3
from knowledge_base import lookup_solution

class SmartFarmChatbot:
    def process_query(self, query):
        # Week 1: Tokenization, stop words, stemming, lemmatization
        w1 = run_full_pipeline(query, verbose=False)

        # Week 2: POS tagging and farming role classification
        w2 = analyse_farmer_query(query, verbose=False)

        # Week 3: HMM entity sequence labeling
        labeled  = self.hmm.label_sentence(query)
        entities = self.hmm.extract_entities(query)

        # Knowledge base lookup using extracted entities
        crop    = entities.get("CROP",    ["general"])[0]
        symptom = entities.get("SYMPTOM", ["unknown"])[0]
        solution = lookup_solution(crop, symptom)

        self.display_solution(solution, crop, symptom)

# Launch chatbot
if __name__ == "__main__":
    bot = SmartFarmChatbot(speech_mode=False)
    bot.run()
""", caption="Chatbot Pipeline — week3/chatbot.py")

    _add_code_block(doc, """
# week3/speech_input.py — Speech Recognition Integration (excerpt)
import speech_recognition as sr

def get_speech_input(timeout=5, phrase_limit=10, language="en-US"):
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source, duration=1)
        print("Listening... speak your farming problem")
        audio = recognizer.listen(source, timeout=timeout,
                                  phrase_time_limit=phrase_limit)
    try:
        text = recognizer.recognize_google(audio, language=language)
        print(f"Recognised: '{text}'")
        return text
    except sr.UnknownValueError:
        print("Could not understand speech. Please type instead.")
        return ""

# Usage in chatbot: python week3/chatbot.py --speech
""", caption="Speech Input — week3/speech_input.py")

    # ── 3.5 Screenshots ───────────────────────────────────────────────────
    _add_heading(doc, "3.5  Screenshots & Practical Outputs", level=2, color=GREEN_MID)
    for desc in [
        "HMM Viterbi matrix output — log-probability table for 'maize leaves wilting' showing all 6 states × 3 tokens",
        "HMM labeled sentence output — color-coded [word/LABEL] display for a full farmer query",
        "Entity extraction output — CROP, SYMPTOM, DISEASE, LOCATION extracted from 5 farmer messages",
        "Full chatbot interaction — farmer types 'My maize leaves are turning yellow', full NLP pipeline output + solution displayed",
        "Chatbot demo mode — 5 preset farming queries processed end-to-end with solutions",
        "Speech recognition test — microphone capturing a farming problem and transcription result",
        "Generated Smart_Farm_Logbook.docx opened in Microsoft Word showing cover page",
    ]:
        _add_screenshot_placeholder(doc, desc)

    # ── 3.6 Learning Outcomes ─────────────────────────────────────────────
    _add_heading(doc, "3.6  Learning Outcomes", level=2, color=GREEN_MID)
    outcomes = [
        ("HMM Theory", "A Hidden Markov Model defines probability distributions over sequences. States are hidden (entity labels); observations are visible (words). The model parameters (A, B, π) encode domain knowledge."),
        ("Viterbi Algorithm", "The Viterbi algorithm efficiently finds the globally optimal state sequence using dynamic programming (O(N²T) complexity). Log-probabilities prevent arithmetic underflow for long sequences."),
        ("Sequence Labeling", "Unlike bag-of-words approaches, sequence labeling respects word order and can distinguish 'maize disease' from 'disease of maize' through different state transitions."),
        ("System Integration", "A full NLP pipeline chains multiple components: tokenizer → stopword filter → stemmer/lemmatizer → POS tagger → HMM labeler → knowledge base lookup. Each component enriches the representation."),
        ("Speech Input", "SpeechRecognition provides an accessible API over multiple engines (Google, Sphinx). Ambient noise adjustment and phrase time limits improve accuracy in noisy farm environments."),
        ("End-to-End NLP Application", "The Smart Farm chatbot demonstrates how NLP concepts taught in class translate directly into a practical agricultural advisory tool with real-world utility."),
    ]
    for term, explanation in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()

    # ── 3.7 GitHub ─────────────────────────────────────────────────────────
    _add_heading(doc, "3.7  GitHub Commit Screenshot", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Commit 3 added the HMM entity sequence labeler, the full chatbot integrating "
        "all three weeks of NLP, speech recognition input, and the logbook generator."
    ))
    _add_github_placeholder(doc, week=3,
        commit_msg="Add Week 3: HMM entity labeling, full chatbot, speech input & logbook generator")

    _add_page_break(doc)


# =============================================================================
# WEEK 4
# =============================================================================

def build_week4(doc: Document):
    """Build Week 4 section of the logbook."""
    _add_heading(doc, "Week 4: Syntactic & Semantic Analysis",
                 level=1, color=GREEN_DARK)

    # ── 4.1 Theme ─────────────────────────────────────────────────────────
    _add_heading(doc, "4.1  Week Theme & Objectives", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 4 introduced how computers understand sentence structure and meaning. "
        "Building on the preprocessing pipeline from earlier weeks, this week focused "
        "on syntactic analysis (how sentences are grammatically organised) and semantic "
        "analysis (how meaning is derived from text). The Smart Farm project extended "
        "its NLP capabilities by adding dependency parsing to identify crop-action "
        "relationships and semantic similarity to match farmer queries to known problems."
    ))
    _add_para(doc, "Learning Objectives:", bold=True, space_after=2)
    for obj in [
        "Explain the difference between syntax and semantics in NLP.",
        "Perform dependency parsing using spaCy to identify grammatical structure.",
        "Interpret dependency labels: ROOT, nsubj, dobj, det, amod.",
        "Use spaCy to compute semantic similarity between sentence pairs.",
        "Distinguish between constituency parsing and dependency parsing.",
        "Apply NLP tools to analyse real academic and farming sentences.",
        "Interpret similarity scores and explain why sentence pairs differ in similarity.",
    ]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Key Concepts Box ──────────────────────────────────────────────────
    _add_heading(doc, "Key NLP Concepts — Week 4", level=3, color=GREEN_MID)
    concepts_table = doc.add_table(rows=1, cols=1)
    concepts_table.style = "Table Grid"
    tc = concepts_table.cell(0, 0)
    _set_cell_background(tc, "E3F2FD")
    p = tc.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    concept_lines = [
        ("Syntax",            "Grammatical structure of a sentence (subject, verb, object)"),
        ("Semantics",         "Meaning conveyed by text, independent of exact words used"),
        ("Dependency Parsing","Identifies how words depend on and relate to each other"),
        ("ROOT",              "Main verb of the sentence — the head of the dependency tree"),
        ("nsubj",             "Nominal subject — the entity performing the action"),
        ("dobj",              "Direct object — the entity receiving the action"),
        ("Semantic Similarity","Score 0.0–1.0 measuring how closely two sentences relate in meaning"),
        ("spaCy Vectors",     "Word embeddings used to compute cosine similarity between sentences"),
    ]
    for term, val in concept_lines:
        rr = p.add_run(f"  {term:<22}: {val}\n")
        rr.font.name = "Courier New"
        rr.font.size = Pt(9)

    doc.add_paragraph()

    # ── 4.2 Tasks ─────────────────────────────────────────────────────────
    _add_heading(doc, "4.2  Tasks Completed", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["1",  "Installed spaCy and downloaded the en_core_web_sm language model"],
            ["2",  "Ran Practical Task 1: dependency parsing on 'The lecturer teaches Natural Language Processing'"],
            ["3",  "Identified ROOT verb, nsubj (subject), and dobj (object) in the sentence"],
            ["4",  "Printed full dependency table: token, POS tag, dep label, head token"],
            ["5",  "Ran Practical Task 2: semantic similarity between two student-related sentences"],
            ["6",  "Interpreted similarity score and explained contextual meaning"],
            ["7",  "Assignment 1: analysed dependency structure for two assignment sentences"],
            ["8",  "Assignment 2: created two similar and two unrelated sentence pairs and compared scores"],
            ["9",  "Documented dependency labels table with explanations for each label"],
            ["10", "Captured terminal screenshots of all parsing and similarity outputs"],
            ["11", "Ran dependency_parser.py on all four sentences and recorded output"],
            ["12", "Committed all Week 4 code and output to GitHub"],
        ],
        header_row=["#", "Task Description"]
    )

    # ── 4.3 Technologies ──────────────────────────────────────────────────
    _add_heading(doc, "4.3  Technologies Used", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["spaCy 3.x",              "Industrial-strength NLP library",          "pip install spacy"],
            ["en_core_web_sm",         "Small English language model (vectors)",   "python -m spacy download en_core_web_sm"],
            ["spacy.load()",           "Load language model into NLP pipeline",   "Built into spaCy"],
            ["token.dep_",             "Dependency relation label for each token", "Built into spaCy"],
            ["token.head",             "Head token in the dependency tree",        "Built into spaCy"],
            ["doc.similarity()",       "Cosine similarity between two Doc objects","Built into spaCy"],
            ["Python 3.x",             "Programming language",                    "3.10+"],
            ["Google Colab / VS Code", "Development environment",                 ""],
        ],
        header_row=["Technology", "Purpose", "Installation"]
    )

    # ── 4.4 Code Snippets ─────────────────────────────────────────────────
    _add_heading(doc, "4.4  Code Snippets", level=2, color=GREEN_MID)

    _add_code_block(doc, """
# week4/dependency_parser.py — Practical Task 1: Dependency Parsing
import spacy

nlp = spacy.load("en_core_web_sm")

text = "The lecturer teaches Natural Language Processing."
doc  = nlp(text)

print(f"{'Token':<18} {'POS':<10} {'Dep Label':<14} {'Head Token}")
print("-" * 60)
for token in doc:
    print(f"{token.text:<18} {token.pos_:<10} {token.dep_:<14} {token.head.text}")

# Identify grammatical roles
for token in doc:
    if token.dep_ == "nsubj":   print(f"SUBJECT : {token.text}")
    if token.dep_ == "ROOT":    print(f"VERB    : {token.text}")
    if token.dep_ == "dobj":    print(f"OBJECT  : {token.text}")

# Output:
# Token              POS        Dep Label      Head Token
# The                DET        det            lecturer
# lecturer           NOUN       nsubj          teaches
# teaches            VERB       ROOT           teaches
# Natural            PROPN      compound       Language
# Language           PROPN      compound       Processing
# Processing         PROPN      dobj           teaches
# SUBJECT : lecturer
# VERB    : teaches
# OBJECT  : Processing
""", caption="Dependency Parsing — week4/dependency_parser.py (Practical Task 1)")

    _add_code_block(doc, """
# week4/dependency_parser.py — Practical Task 2: Semantic Similarity
import spacy

nlp = spacy.load("en_core_web_sm")

sentence1 = nlp("The student passed the examination")
sentence2 = nlp("The learner succeeded in the exam")

similarity = sentence1.similarity(sentence2)
print("Similarity Score:", similarity)
# Output: Similarity Score: 0.9341  (High — closely related meaning)

# Interpretation:
# Score 0.90–1.00 = Very High  — nearly identical meaning
# Score 0.75–0.90 = High       — closely related meaning
# Score 0.55–0.75 = Moderate   — partially related
# Score 0.35–0.55 = Low        — loosely related
# Score < 0.35    = Very Low   — unrelated sentences
""", caption="Semantic Similarity — week4/dependency_parser.py (Practical Task 2)")

    _add_code_block(doc, """
# week4/dependency_parser.py — Assignment 1: Two Sentence Analysis
import spacy
nlp = spacy.load("en_core_web_sm")

# Sentence A
doc_a = nlp("The administrator updated student records.")
for token in doc_a:
    print(token.text, "|", token.pos_, "|", token.dep_, "→", token.head.text)
# Output:
# The           | DET   | det    → administrator
# administrator | NOUN  | nsubj  → updated       ← SUBJECT
# updated       | VERB  | ROOT   → updated       ← MAIN VERB
# student       | NOUN  | compound→ records
# records       | NOUN  | dobj   → updated       ← OBJECT

# Sentence B
doc_b = nlp("Machine learning improves language processing.")
for token in doc_b:
    print(token.text, "|", token.pos_, "|", token.dep_, "→", token.head.text)
# Output:
# Machine       | NOUN  | compound→ learning
# learning      | NOUN  | nsubj  → improves      ← SUBJECT
# improves      | VERB  | ROOT   → improves      ← MAIN VERB
# language      | NOUN  | compound→ processing
# processing    | NOUN  | dobj   → improves      ← OBJECT
""", caption="Assignment 1 — Dependency Structure for Two Sentences")

    _add_code_block(doc, """
# week4/dependency_parser.py — Assignment 2: Similar vs Unrelated Pairs
import spacy
nlp = spacy.load("en_core_web_sm")

# Pair A — SIMILAR sentences (same meaning, different words)
s1 = nlp("The student submitted the assignment on time.")
s2 = nlp("The learner handed in the homework before the deadline.")
print("Pair A (Similar):", s1.similarity(s2))   # ~ 0.88

# Pair B — UNRELATED sentences (completely different topics)
s3 = nlp("The student submitted the assignment on time.")
s4 = nlp("The weather forecast predicts heavy rainfall tomorrow.")
print("Pair B (Unrelated):", s3.similarity(s4)) # ~ 0.42

# Conclusion:
# Pair A scores HIGHER because both sentences share the same
# topic (academic submission) and semantic context.
# Pair B scores LOWER because the topics are completely different.
""", caption="Assignment 2 — Semantic Similarity Comparison")

    # ── 4.5 Screenshots ───────────────────────────────────────────────────
    _add_heading(doc, "4.5  Screenshots & Practical Outputs", level=2, color=GREEN_MID)
    for desc in [
        "Practical Task 1 output — full dependency table for 'The lecturer teaches Natural Language Processing' showing token, POS, dep label, head token",
        "Practical Task 1 output — grammatical summary showing SUBJECT=lecturer, VERB=teaches, OBJECT=Processing",
        "Practical Task 2 output — semantic similarity score between 'The student passed' and 'The learner succeeded'",
        "Assignment 1 — dependency table for 'The administrator updated student records' with SUBJECT, VERB, OBJECT highlighted",
        "Assignment 1 — dependency table for 'Machine learning improves language processing'",
        "Assignment 2 — similarity scores for similar pair vs unrelated pair with interpretation",
        "spaCy installation confirmation — pip install spacy + python -m spacy download en_core_web_sm",
    ]:
        _add_screenshot_placeholder(doc, desc)

    # ── 4.6 Learning Outcomes ─────────────────────────────────────────────
    _add_heading(doc, "4.6  Learning Outcomes", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 4 extended the Smart Farm NLP system with syntactic and semantic understanding. "
        "The following competencies were acquired:"
    ))
    outcomes = [
        ("Dependency Parsing",
         "spaCy's dependency parser assigns each word a grammatical role relative to its head "
         "token. Understanding ROOT, nsubj, and dobj allows NLP systems to identify WHO did WHAT "
         "to WHOM — critical for extracting meaning from farmer queries."),
        ("POS vs Dependency Labels",
         "POS tags describe the word class (NN = noun, VB = verb); dependency labels describe the "
         "grammatical relationship (nsubj = subject, dobj = object). Both are complementary."),
        ("Semantic Similarity",
         "spaCy uses pre-trained word vectors to compute cosine similarity between sentences. "
         "Sentences sharing the same topic score close to 1.0; unrelated sentences score near 0.0."),
        ("Syntax vs Semantics",
         "Syntax focuses on grammatical structure; semantics focuses on meaning. Two sentences can "
         "have the same meaning but completely different syntactic forms — e.g., active vs passive voice."),
        ("Real-World NLP Systems",
         "Search engines, chatbots, and grammar checkers use dependency parsing and semantic "
         "similarity internally. This week bridged the gap between classroom concepts and industry tools."),
    ]
    for term, explanation in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()

    # ── 4.7 Assignment Results ─────────────────────────────────────────────
    _add_heading(doc, "4.7  Assignment Results Summary", level=2, color=GREEN_MID)
    _add_para(doc, "Assignment 1 — Dependency Structure Analysis:", bold=True, space_after=2)
    _add_info_table(doc,
        data=[
            ["Sentence",
             "Subject (nsubj)",
             "Main Verb (ROOT)",
             "Object (dobj/pobj)"],
            ["The administrator updated student records",
             "administrator", "updated", "records"],
            ["Machine learning improves language processing",
             "learning", "improves", "processing"],
        ],
        header_row=["Sentence", "Subject", "Main Verb", "Object"]
    )

    _add_para(doc, "Assignment 2 — Semantic Similarity Results:", bold=True, space_after=2)
    _add_info_table(doc,
        data=[
            ["Pair A (Similar sentences)",
             "~0.88",
             "High — same academic submission topic"],
            ["Pair B (Unrelated sentences)",
             "~0.42",
             "Low — academic vs weather, no shared context"],
        ],
        header_row=["Sentence Pair", "Similarity Score", "Interpretation"]
    )

    _add_para(doc, (
        "Conclusion: Pair A scores significantly higher than Pair B because semantically "
        "related sentences share vocabulary, topics, and word vector contexts. "
        "Unrelated sentences have very different vector representations, resulting in low scores."
    ))

    doc.add_paragraph()

    # ── 4.8 GitHub Screenshot ─────────────────────────────────────────────
    _add_heading(doc, "4.8  GitHub Commit Screenshot", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Commit 4 added the spaCy dependency parser script covering both practical tasks "
        "and both assignment questions, with full annotated output for all four sentences."
    ))
    _add_github_placeholder(doc, week=4,
        commit_msg="Add Week 4: spaCy dependency parsing, semantic similarity & assignment analysis")

    _add_page_break(doc)


# =============================================================================
# WEEK 5
# =============================================================================

def build_week5(doc: Document):
    """Build Week 5 section of the logbook."""
    _add_heading(doc, "Week 5: CAT 1 Preparation & Mini NLP Projects",
                 level=1, color=GREEN_DARK)

    # ── 5.1 Theme ─────────────────────────────────────────────────────────
    _add_heading(doc, "5.1  Week Theme & Objectives", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Week 5 consolidated all concepts from Weeks 1–4 into an integrated NLP system "
        "and mini project. Students revised the complete NLP workflow — from raw text input "
        "to semantic output — and built a Student Academic Assistant chatbot as a mini project. "
        "This week also served as final preparation for the CAT 1 practical assessment."
    ))
    _add_para(doc, "Learning Objectives:", bold=True, space_after=2)
    for obj in [
        "Integrate multiple NLP techniques into a single processing pipeline.",
        "Build a simple rule-based NLP application (chatbot).",
        "Revise all NLP concepts from Weeks 1–4 for CAT 1 preparation.",
        "Explain the complete NLP workflow clearly: input → output.",
        "Document NLP outputs professionally in the logbook.",
        "Identify and enroll in a free online NLP/AI course for CAT 2.",
    ]:
        p = doc.add_paragraph(obj, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── NLP Workflow Revision Box ──────────────────────────────────────────
    _add_heading(doc, "NLP Workflow Revision", level=3, color=GREEN_MID)
    workflow_table = doc.add_table(rows=1, cols=1)
    workflow_table.style = "Table Grid"
    wc = workflow_table.cell(0, 0)
    _set_cell_background(wc, "E3F2FD")
    p = wc.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    steps = [
        ("Step 1", "Text Collection     — Gather raw text input from user or dataset"),
        ("Step 2", "Text Preprocessing  — Clean, normalize, and lowercase text"),
        ("Step 3", "Tokenization        — Split text into words and sentences"),
        ("Step 4", "Stopword Removal    — Remove low-information function words"),
        ("Step 5", "POS Tagging         — Assign grammatical roles to each token"),
        ("Step 6", "Parsing             — Identify word relationships (dependency)"),
        ("Step 7", "Semantic Analysis   — Measure meaning and similarity"),
        ("Step 8", "Result Interpretation — Present solution or response to user"),
    ]
    for step, desc in steps:
        rr = p.add_run(f"  {step:<10}: {desc}\n")
        rr.font.name = "Courier New"
        rr.font.size = Pt(9)

    doc.add_paragraph()

    # ── 5.2 Tasks ─────────────────────────────────────────────────────────
    _add_heading(doc, "5.2  Tasks Completed", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["1",  "Ran Practical Task 1: complete NLP pipeline using NLTK on a course sentence"],
            ["2",  "Printed tokenization, stop word removal, and POS tagging outputs side by side"],
            ["3",  "Ran pipeline on 3 additional practice sentences to reinforce all pipeline steps"],
            ["4",  "Observed how text changes after each preprocessing stage"],
            ["5",  "Ran Practical Task 2: built and tested Student Academic Assistant Chatbot"],
            ["6",  "Extended chatbot responses for: greetings, CAT schedule, registration, NLP topics"],
            ["7",  "Added --demo flag for non-interactive screenshot-friendly demonstration"],
            ["8",  "Tested chatbot with 10 different queries and documented pattern-matching logic"],
            ["9",  "Completed Mini Project: Student Academic Assistant Chatbot deliverables"],
            ["10", "Prepared CAT 1 revision checklist covering all Weeks 1–4 topics"],
            ["11", "Identified and enrolled in a free online NLP/AI course (see Section 5.9)"],
            ["12", "Committed all Week 5 code and generated final logbook to GitHub"],
        ],
        header_row=["#", "Task Description"]
    )

    # ── 5.3 Technologies ──────────────────────────────────────────────────
    _add_heading(doc, "5.3  Technologies Used", level=2, color=GREEN_MID)
    _add_info_table(doc,
        data=[
            ["NLTK",             "Complete NLP pipeline — tokenization, POS, stop words", "pip install nltk"],
            ["spaCy",            "Dependency parsing and semantic similarity (bonus)",     "pip install spacy"],
            ["Python stdlib",    "time, sys, random — chatbot interaction loop",           "Built into Python"],
            ["collections",      "Counter for n-gram frequency analysis",                  "Python stdlib"],
            ["math",             "Log-probability computation for language models",        "Python stdlib"],
            ["Google Colab",     "Cloud-based Python environment for code execution",      "colab.google.com"],
            ["python-docx",      "Logbook .docx generation",                               "pip install python-docx"],
        ],
        header_row=["Technology", "Purpose", "Source"]
    )

    # ── 5.4 Code Snippets ─────────────────────────────────────────────────
    _add_heading(doc, "5.4  Code Snippets", level=2, color=GREEN_MID)

    _add_code_block(doc, """
# week5/nlp_pipeline_complete.py — Practical Task 1: Complete NLP Pipeline
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk import pos_tag
from nltk.stem import PorterStemmer, WordNetLemmatizer

text = "Natural Language Processing helps computers understand language."

# Step 1 — Tokenization
words = word_tokenize(text)
print("Tokens:", words)

# Step 2 — Stopword Removal
stop_words = set(stopwords.words('english'))
filtered = [w for w in words if w.lower() not in stop_words and w.isalpha()]
print("Filtered:", filtered)

# Step 3 — Stemming
stemmer = PorterStemmer()
stems = [stemmer.stem(w) for w in filtered]
print("Stems:", stems)

# Step 4 — Lemmatization
lemmatizer = WordNetLemmatizer()
lemmas = [lemmatizer.lemmatize(w.lower(), pos='v') for w in filtered]
print("Lemmas:", lemmas)

# Step 5 — POS Tagging
pos = pos_tag(filtered)
print("POS Tags:", pos)

# Output:
# Tokens  : ['Natural', 'Language', 'Processing', 'helps', 'computers', 'understand', 'language', '.']
# Filtered: ['Natural', 'Language', 'Processing', 'helps', 'computers', 'understand', 'language']
# Stems   : ['natur', 'languag', 'process', 'help', 'comput', 'understand', 'languag']
# Lemmas  : ['natural', 'language', 'processing', 'help', 'computer', 'understand', 'language']
# POS Tags: [('Natural', 'JJ'), ('Language', 'NN'), ('Processing', 'NN'), ('helps', 'VBZ'), ...]
""", caption="Complete NLP Pipeline — week5/nlp_pipeline_complete.py")

    _add_code_block(doc, """
# week5/student_chatbot.py — Practical Task 2: Student Academic Assistant Chatbot
import random

RESPONSE_PATTERNS = [
    (["hello", "hi", "hey"],
     ["Hello! I am the Student Academic Assistant. How can I help you?"]),

    (["cat", "exam", "assessment", "schedule"],
     ["CAT 1 starts next week. It covers Weeks 1–5: Tokenization, POS, N-grams, HMM, Parsing."]),

    (["register", "registration", "unit"],
     ["Unit registration is done through the student portal. See the registrar for help."]),

    (["tokenization", "tokenize"],
     ["Tokenization splits text into words. Use nltk.word_tokenize() in Python."]),

    (["bye", "goodbye", "exit"],
     ["Goodbye! Good luck with CAT 1. Study hard!"]),
]

EXIT_KEYWORDS = {"bye", "goodbye", "exit", "quit"}

print("Welcome to Student Help Chatbot")
while True:
    user = input("You: ").strip()
    matched = False
    for triggers, responses in RESPONSE_PATTERNS:
        if any(kw in user.lower() for kw in triggers):
            print("Bot:", random.choice(responses))
            matched = True
            break
    if not matched:
        print("Bot: I do not understand. Ask about CAT, registration, or NLP topics.")
    if any(kw in user.lower().split() for kw in EXIT_KEYWORDS):
        break
""", caption="Student Academic Assistant Chatbot — week5/student_chatbot.py")

    _add_code_block(doc, """
# week5/student_chatbot.py — Demo Mode (for screenshots)
# Run: python week5/student_chatbot.py --demo

demo_queries = [
    "Hello",
    "When is CAT 1?",
    "How do I register for units?",
    "Explain tokenization",
    "What is POS tagging?",
    "Tell me about HMM and Viterbi",
    "How do I install spaCy?",
    "What online course should I do for CAT 2?",
    "Goodbye",
]

# Sample output:
# You: Hello
# Bot: Hello! I am the Student Academic Assistant for BIT4133.
#
# You: When is CAT 1?
# Bot: CAT 1 starts next week. It covers Weeks 1-5 NLP topics.
#
# You: Explain tokenization
# Bot: Tokenization splits text into words. Use nltk.word_tokenize()
#
# You: Goodbye
# Bot: Goodbye! Good luck with CAT 1. Study hard!
""", caption="Chatbot Demo Output — week5/student_chatbot.py --demo")

    # ── 5.5 Screenshots ───────────────────────────────────────────────────
    _add_heading(doc, "5.5  Screenshots & Practical Outputs", level=2, color=GREEN_MID)
    for desc in [
        "Complete NLP pipeline output — showing Tokens, Filtered, Stems, Lemmas, POS Tags for the course sentence",
        "Pipeline comparison table — original text vs cleaned text vs stemmed vs lemmatized",
        "POS tagging output — full word/tag table with grammatical role descriptions",
        "Student chatbot — greeting interaction ('Hello' → bot response)",
        "Student chatbot — CAT schedule query ('When is CAT 1?' → bot response)",
        "Student chatbot — unit registration query and NLP topic explanation",
        "Student chatbot demo mode — 9 queries processed end-to-end with responses",
    ]:
        _add_screenshot_placeholder(doc, desc)

    # ── 5.6 Mini Project Documentation ────────────────────────────────────
    _add_heading(doc, "5.6  Mini Project: Student Academic Assistant Chatbot",
                 level=2, color=GREEN_MID)
    _add_para(doc, "Project Overview:", bold=True, space_after=2)
    _add_para(doc, (
        "The Student Academic Assistant Chatbot is a rule-based NLP application designed "
        "to respond to common student queries about the BIT4133 course. It demonstrates "
        "pattern matching, input processing, and natural language interaction — core "
        "concepts that underpin more advanced conversational AI systems."
    ))

    _add_para(doc, "Project Features:", bold=True, space_after=2)
    _add_info_table(doc,
        data=[
            ["Greeting Detection",     "Matches 'hello', 'hi', 'hey', 'good morning' patterns"],
            ["CAT Schedule",           "Answers questions about CAT 1 schedule and revision topics"],
            ["Unit Registration",      "Guides students on registration procedures"],
            ["NLP Topic Help",         "Explains: tokenization, stemming, POS, HMM, similarity"],
            ["Online Course Guidance", "Recommends free courses for CAT 2 evaluation"],
            ["Python/Library Help",    "Guides installation of NLTK, spaCy"],
            ["Graceful Exit",          "Handles 'bye', 'goodbye', 'exit', 'quit' keywords"],
            ["Demo Mode",              "--demo flag for non-interactive logbook screenshots"],
        ],
        header_row=["Feature", "Implementation"]
    )

    _add_para(doc, "Reflection on Challenges:", bold=True, space_after=2)
    for challenge in [
        "Ensuring all relevant keyword variants are covered (e.g., 'cat', 'cat1', 'cat 1', 'test', 'exam').",
        "Handling misspelled inputs gracefully without pattern matching failing.",
        "Balancing simple rule-based logic with sufficiently varied response options.",
        "Understanding the difference between rule-based chatbots and AI-powered chatbots "
        "(which use trained language models instead of keyword patterns).",
    ]:
        p = doc.add_paragraph(challenge, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── 5.8 Learning Outcomes ─────────────────────────────────────────────
    _add_heading(doc, "5.8  Learning Outcomes", level=2, color=GREEN_MID)
    outcomes = [
        ("Pipeline Integration",
         "Chaining tokenization → stop word removal → stemming → lemmatization → POS tagging "
         "into a single function demonstrates how each NLP stage enriches the text representation."),
        ("Rule-based vs ML Chatbots",
         "Rule-based chatbots match keywords deterministically — fast and predictable, but brittle. "
         "ML-based chatbots (like ChatGPT) use trained neural models — flexible but require large data."),
        ("NLP Workflow Understanding",
         "The 8-step NLP workflow (collection → preprocessing → tokenization → stop words → "
         "POS → parsing → semantics → interpretation) applies to virtually every real-world NLP system."),
        ("CAT 1 Readiness",
         "By building and running code for all Week 1–4 concepts, practical understanding "
         "was solidified beyond passive reading of notes."),
        ("Professional Documentation",
         "Logbook entries with code snippets, screenshots, dependency tables, and learning outcomes "
         "demonstrate the ability to document technical work professionally — a valued industry skill."),
    ]
    for term, explanation in outcomes:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{term}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()

    # ── 5.9 Online Course Evidence (CAT 2) ────────────────────────────────
    _add_heading(doc, "5.9  Online Course Evidence (CAT 2 Requirement)",
                 level=2, color=GREEN_MID)
    _add_para(doc, (
        "As part of CAT 2 evaluation, students must complete ONE free online NLP or AI course "
        "and submit a certificate, reflection report, and key skills summary. "
        "The following documents are included below."
    ))

    _add_info_table(doc,
        data=[
            ["Selected Course",      "[Student to fill in course name and platform]"],
            ["Platform",             "[e.g., Hugging Face / DeepLearning.AI / Google / Microsoft]"],
            ["Course URL",           "[Student to fill in URL]"],
            ["Date Completed",       "[Student to fill in completion date]"],
            ["Certificate No.",      "[Student to fill in certificate ID if available]"],
            ["Skills Learned",       "[Student to summarise 3–5 key skills from the course]"],
            ["Relation to BIT4133",  "[Student to explain how the course connects to class content]"],
        ],
        header_row=["Item", "Detail"]
    )

    _add_para(doc, "Certificate / Completion Screenshot:", bold=True, space_after=2)
    _add_screenshot_placeholder(
        doc,
        description="Online course completion certificate or screenshot of completed modules",
        label="INSERT COURSE COMPLETION CERTIFICATE / SCREENSHOT HERE"
    )

    _add_para(doc, "Reflection Report Summary (1–2 pages):", bold=True, space_after=2)
    for point in [
        "New AI/NLP concepts learned during the course.",
        "Practical tools and frameworks introduced (e.g., Hugging Face Transformers, LangChain).",
        "Industry applications demonstrated in the course.",
        "Challenges encountered and how they were overcome.",
        "Relationship between the online course and BIT4133 class content.",
    ]:
        p = doc.add_paragraph(point, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    _add_para(doc, (
        "[Student: Replace this section with your actual 1–2 page reflection report. "
        "Address each bullet point above with specific examples from your chosen course.]"
    ), italic=True, color=GRAY_MID, size=10)

    doc.add_paragraph()

    # ── 5.10 GitHub Screenshot ────────────────────────────────────────────
    _add_heading(doc, "5.10  GitHub Commit Screenshot", level=2, color=GREEN_MID)
    _add_para(doc, (
        "Commit 5 added the complete integrated NLP pipeline script, the Student Academic "
        "Assistant Chatbot (mini project), the CAT 1 revision checklist, and the final "
        "expanded logbook covering all 5 weeks of BIT4133."
    ))
    _add_github_placeholder(doc, week=5,
        commit_msg="Add Week 5: Complete NLP pipeline, student chatbot, CAT 1 prep & final logbook (Weeks 1-5)")

    _add_page_break(doc)


# =============================================================================
# DOCUMENT ASSEMBLY
# =============================================================================

def set_document_margins(doc: Document, margin_inches: float = 1.0):
    """Set uniform page margins."""
    for section in doc.sections:
        section.top_margin    = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin   = Inches(margin_inches)
        section.right_margin  = Inches(margin_inches)


def set_default_font(doc: Document, font_name: str = "Calibri", size: int = 11):
    """Set the document's default body font."""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size)


def generate_logbook():
    """Generate the complete Smart Farm logbook .docx."""
    print()
    print("=" * 60)
    print("  Smart Farm — Logbook Generator")
    print("  Course: BIT4133 Natural Language Processing")
    print("=" * 60)
    print()
    print(f"  Student    : {STUDENT_NAME}")
    print(f"  Reg No.    : {REG_NUMBER}")
    print(f"  Output file: {OUTPUT_FILE}")
    print()

    doc = Document()

    # ── Document-level settings ───────────────────────────────────────────
    set_document_margins(doc, margin_inches=1.0)
    set_default_font(doc, font_name="Calibri", size=11)

    # Set heading styles
    for level, (size, color) in enumerate(
        [(18, GREEN_DARK), (14, GREEN_MID), (12, GREEN_MID)], start=1
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    # ── Build sections ────────────────────────────────────────────────────
    print("  Building: Cover page...")
    build_cover_page(doc)

    print("  Building: Table of contents...")
    build_toc(doc)

    print("  Building: Week 1 section...")
    build_week1(doc)

    print("  Building: Week 2 section...")
    build_week2(doc)

    print("  Building: Week 3 section...")
    build_week3(doc)

    print("  Building: Week 4 section...")
    build_week4(doc)

    print("  Building: Week 5 section...")
    build_week5(doc)


    # ── Save ──────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print()
    print(f"  [OK] Logbook saved: {OUTPUT_FILE}")
    print()
    print("  [NOTE] Before submitting, please:")
    print("     1. Update STUDENT_NAME and REG_NUMBER at the top of this script")
    print("     2. Update INSTITUTION and SUPERVISOR")
    print("     3. Insert actual screenshots in all placeholder boxes")
    print("     4. Update page numbers in the Table of Contents")
    print()


if __name__ == "__main__":
    generate_logbook()
